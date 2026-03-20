"""
Translation endpoint — English ↔ Malayalam legal text translation.

Routes
------
POST /api/v1/translate/text
    Translate a plain-text payload.

POST /api/v1/translate/text/stream
    Translate a plain-text payload via SSE streaming.

POST /api/v1/translate/document
    Translate an uploaded file (PDF / DOCX / TXT).

POST /api/v1/translate/export
    Export translated text as PDF or DOCX file download.
"""
from __future__ import annotations

import io
import json
import logging
from typing import List, Literal, Optional

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile, status
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field

from app.api.v1.deps import get_current_user
from app.db.models import User
from app.services.translation.llm_translate_service import llm_translate_service
from app.services.translation.document_translate_service import document_translate_service

logger = logging.getLogger(__name__)

router = APIRouter()

Direction = Literal["en_to_ml", "ml_to_en"]
DirectionInput = Literal["en_to_ml", "ml_to_en", "auto"]

_ALLOWED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",
    "text/plain",
}

_MAX_TEXT_CHARS = 15_000   # ~3 000 words
_MAX_FILE_BYTES = 10 * 1024 * 1024  # 10 MB


# ── Direction auto-detection (Task 3) ──────────────────────────────────────

def detect_direction(text: str) -> Direction:
    """
    Auto-detect the translation direction from script composition.

    Counts alphabetic characters in the Malayalam Unicode block (U+0D00–U+0D7F).
    If more than 30 % of all alphabetic characters are Malayalam, the text is
    assumed to be Malayalam and the direction is "ml_to_en"; otherwise "en_to_ml".
    """
    total = 0
    malayalam = 0
    for ch in text:
        if ch.isalpha():
            total += 1
            if 0x0D00 <= ord(ch) <= 0x0D7F:
                malayalam += 1
    if total == 0:
        return "en_to_ml"
    return "ml_to_en" if (malayalam / total) > 0.30 else "en_to_ml"


# ── Request / response schemas ─────────────────────────────────────────────

class GlossaryTerm(BaseModel):
    source: str
    target: str


class TranslateTextRequest(BaseModel):
    text: str = Field(..., min_length=1, max_length=_MAX_TEXT_CHARS)
    direction: DirectionInput = "en_to_ml"


class TranslateTextResponse(BaseModel):
    translated: str
    direction: str
    glossary_hits: int
    warnings: List[str]
    char_count: int
    glossary_terms: List[GlossaryTerm] = []


class TranslateDocumentResponse(BaseModel):
    translated: str
    direction: str
    filename: str
    mime_type: str
    chunks: int
    glossary_hits: int
    warnings: List[str]
    char_count: int


# ── Endpoints ──────────────────────────────────────────────────────────────


@router.post(
    "/text",
    response_model=TranslateTextResponse,
    summary="Translate legal text",
    description=(
        "Translate a plain-text legal document between English and Malayalam. "
        "Pass direction='auto' to let the service detect the language automatically. "
        "Protected entities (case numbers, acts, dates, Latin phrases) are preserved verbatim."
    ),
)
def translate_text(
    payload: TranslateTextRequest,
    current_user: User = Depends(get_current_user),
) -> TranslateTextResponse:
    """Translate a plain-text payload."""
    # Resolve direction
    direction: Direction = (
        detect_direction(payload.text)
        if payload.direction == "auto"
        else payload.direction  # type: ignore[assignment]
    )

    try:
        result = llm_translate_service.translate_text(payload.text, direction)
    except RuntimeError as exc:
        logger.error("translate/text failed: %s", exc)
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=f"Translation service error: {exc}",
        )

    glossary_terms = [
        GlossaryTerm(source=g["source"], target=g["target"])
        for g in result.get("glossary_terms", [])
    ]

    return TranslateTextResponse(
        translated=result["translated"],
        direction=direction,
        glossary_hits=result["glossary_hits"],
        warnings=result["warnings"],
        char_count=len(payload.text),
        glossary_terms=glossary_terms,
    )


@router.post(
    "/text/stream",
    summary="Stream-translate legal text (SSE)",
    description=(
        "Stream a translation using Server-Sent Events. "
        "Raw token chunks are emitted as they arrive; "
        "the final 'done' event contains the fully restored clean translation."
    ),
)
def translate_text_stream(
    payload: TranslateTextRequest,
    current_user: User = Depends(get_current_user),
) -> StreamingResponse:
    """Stream the translation of a plain-text payload via SSE."""
    direction: Direction = (
        detect_direction(payload.text)
        if payload.direction == "auto"
        else payload.direction  # type: ignore[assignment]
    )

    def generate():
        try:
            for item in llm_translate_service.stream_translate_text(
                payload.text, direction
            ):
                if isinstance(item, dict):
                    # Final done event — add char_count and direction
                    item["char_count"] = len(payload.text)
                    item["direction"] = direction
                    yield f"data: {json.dumps(item)}\n\n"
                else:
                    # Raw text chunk
                    yield f"data: {json.dumps({'text': item})}\n\n"
        except RuntimeError as exc:
            logger.error("translate/text/stream failed: %s", exc)
            yield f"data: {json.dumps({'error': str(exc)})}\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


@router.post(
    "/document/stream",
    summary="Translate uploaded document — chunk-by-chunk SSE stream",
    description=(
        "Upload a PDF, DOCX, or TXT file. The backend extracts text (OCR for scanned PDFs), "
        "then translates each chunk and streams it via SSE as soon as it is ready. "
        "Event types: 'extracted' (OCR done, total chunk count known), "
        "'chunk' (one translated chunk), 'done' (final quality metadata), 'error'."
    ),
)
async def translate_document_stream(
    file: UploadFile = File(...),
    direction: Direction = Form("en_to_ml"),
    current_user: User = Depends(get_current_user),
) -> StreamingResponse:
    """Translate an uploaded document with per-chunk SSE streaming."""
    from app.services.translation.document_translate_service import extract_text, chunk_text
    from app.services.translation.protect_service import protect_service
    from app.services.translation.glossary_service import glossary_service

    content_type = (file.content_type or "").split(";")[0].strip()
    if content_type not in _ALLOWED_MIME_TYPES:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=f"Unsupported file type '{content_type}'. Allowed: PDF, DOCX, TXT.",
        )

    data = await file.read()
    if len(data) > _MAX_FILE_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File too large ({len(data) // 1024} KB). Maximum is 10 MB.",
        )

    filename = file.filename or "document"

    def generate():
        # ── Phase 1: OCR / text extraction ──────────────────────────────────
        try:
            text = extract_text(data, content_type, ocr_lang="mal+eng")
        except (ValueError, RuntimeError) as exc:
            yield f"data: {json.dumps({'type': 'error', 'message': str(exc)})}\n\n"
            return

        chunks = chunk_text(text)
        yield f"data: {json.dumps({'type': 'extracted', 'total_chunks': len(chunks), 'char_count': len(text)})}\n\n"

        # ── Phase 2: translate chunk by chunk ────────────────────────────────
        total_glossary_hits = 0
        all_warnings: list = []
        translated_parts: list = []

        for idx, chunk in enumerate(chunks):
            protected, pmap = protect_service.protect_text(chunk)
            try:
                translated_protected = llm_translate_service.translate_chunk(
                    protected, direction
                )
            except RuntimeError as exc:
                logger.error("Document stream chunk %d/%d failed: %s", idx + 1, len(chunks), exc)
                translated_parts.append(chunk)
                all_warnings.append(f"Chunk {idx + 1} failed: {exc}")
                yield f"data: {json.dumps({'type': 'chunk', 'index': idx, 'total': len(chunks), 'text': chunk, 'failed': True})}\n\n"
                continue

            restored = protect_service.restore_text(translated_protected, pmap)
            warnings = protect_service.validate_protection(chunk, restored)
            all_warnings.extend(warnings)

            hits = len(glossary_service.find_matches(chunk, direction))
            total_glossary_hits += hits
            translated_parts.append(restored)

            yield f"data: {json.dumps({'type': 'chunk', 'index': idx, 'total': len(chunks), 'text': restored})}\n\n"

        # ── Phase 3: done ─────────────────────────────────────────────────────
        yield f"data: {json.dumps({'type': 'done', 'total_chunks': len(chunks), 'glossary_hits': total_glossary_hits, 'warnings': all_warnings, 'char_count': len(text), 'filename': filename, 'direction': direction})}\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


@router.post(
    "/document",
    response_model=TranslateDocumentResponse,
    summary="Translate uploaded legal document",
    description=(
        "Upload a PDF, DOCX, or TXT file to translate its full text. "
        "Large documents are split into chunks automatically."
    ),
)
async def translate_document(
    file: UploadFile = File(...),
    direction: Direction = Form("en_to_ml"),
    current_user: User = Depends(get_current_user),
) -> TranslateDocumentResponse:
    """Translate an uploaded document."""
    content_type = (file.content_type or "").split(";")[0].strip()
    if content_type not in _ALLOWED_MIME_TYPES:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=(
                f"Unsupported file type '{content_type}'. "
                "Allowed: PDF, DOCX, TXT."
            ),
        )

    data = await file.read()
    if len(data) > _MAX_FILE_BYTES:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File too large ({len(data) // 1024} KB). Maximum is 10 MB.",
        )

    try:
        result = document_translate_service.translate_bytes(data, content_type, direction)
    except ValueError as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail=str(exc)
        )
    except RuntimeError as exc:
        logger.error("translate/document failed: %s", exc)
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=f"Translation service error: {exc}",
        )

    return TranslateDocumentResponse(
        translated=result["translated"],
        direction=direction,
        filename=file.filename or "document",
        mime_type=content_type,
        chunks=result["chunks"],
        glossary_hits=result["glossary_hits"],
        warnings=result["warnings"],
        char_count=result.get("char_count", 0),
    )


# ── Export ─────────────────────────────────────────────────────────────────

ExportFormat = Literal["pdf", "docx"]


class ExportRequest(BaseModel):
    text: str = Field(..., min_length=1)
    title: str = Field(default="Translation", max_length=200)
    direction: Direction = "en_to_ml"
    format: ExportFormat = "pdf"


def _build_pdf(text: str, title: str, direction: str) -> bytes:
    """Generate a PDF using reportlab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    from reportlab.lib.enums import TA_LEFT
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import os

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()

    font_name = "Helvetica"
    font_candidates = [
        "/System/Library/Fonts/Supplemental/NotoSansMalayalam.ttf",
        "/usr/share/fonts/truetype/noto/NotoSansMalayalam-Regular.ttf",
        "/usr/share/fonts/NotoSansMalayalam-Regular.ttf",
    ]
    for candidate in font_candidates:
        if os.path.exists(candidate):
            try:
                pdfmetrics.registerFont(TTFont("NotoMal", candidate))
                font_name = "NotoMal"
            except Exception:
                pass
            break

    title_style = ParagraphStyle(
        "TransTitle",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=14,
        spaceAfter=12,
    )
    body_style = ParagraphStyle(
        "TransBody",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=11,
        leading=18,
        spaceAfter=8,
    )
    meta_style = ParagraphStyle(
        "TransMeta",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=9,
        textColor=(0.5, 0.5, 0.5),
        spaceAfter=20,
    )

    dir_label = "English → Malayalam" if direction == "en_to_ml" else "Malayalam → English"
    story = [
        Paragraph(title, title_style),
        Paragraph(f"Legal Translation  ·  {dir_label}", meta_style),
    ]

    for para in text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        safe = para.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        story.append(Paragraph(safe, body_style))
        story.append(Spacer(1, 4))

    doc.build(story)
    return buf.getvalue()


def _build_docx(text: str, title: str, direction: str) -> bytes:
    """Generate a DOCX using python-docx."""
    import docx

    document = docx.Document()
    document.add_heading(title, level=1)

    dir_label = "English → Malayalam" if direction == "en_to_ml" else "Malayalam → English"
    meta = document.add_paragraph()
    run = meta.add_run(f"Legal Translation  ·  {dir_label}")
    run.italic = True
    run.font.color.rgb = docx.shared.RGBColor(0x88, 0x88, 0x88)

    document.add_paragraph()

    for para in text.split("\n\n"):
        para = para.strip()
        if not para:
            continue
        p = document.add_paragraph(para)
        p.style = document.styles["Normal"]

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


@router.post(
    "/export",
    summary="Export translated text as PDF or DOCX",
    response_class=StreamingResponse,
)
def export_translation(
    payload: ExportRequest,
    current_user: User = Depends(get_current_user),
) -> StreamingResponse:
    """Generate a formatted PDF or DOCX from translated text."""
    try:
        if payload.format == "pdf":
            file_bytes = _build_pdf(payload.text, payload.title, payload.direction)
            media_type = "application/pdf"
            ext = "pdf"
        else:
            file_bytes = _build_docx(payload.text, payload.title, payload.direction)
            media_type = (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            ext = "docx"
    except Exception as exc:
        logger.error("translate/export failed: %s", exc)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Export failed: {exc}",
        )

    safe_title = "".join(
        c if c.isalnum() or c in "-_ " else "_" for c in payload.title
    )
    filename = f"{safe_title}_{payload.direction}.{ext}"

    return StreamingResponse(
        io.BytesIO(file_bytes),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
